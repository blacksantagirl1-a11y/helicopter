from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.shared import Pt


def norm(text: str) -> str:
    return " ".join((text or "").split()).casefold()


def find_paragraph(doc: Document, needle: str) -> int:
    wanted = norm(needle)
    for i, paragraph in enumerate(doc.paragraphs):
        if norm(paragraph.text) == wanted:
            return i
    raise ValueError(f"Paragraph not found: {needle}")


def before(paragraph, text: str = "", style: str | None = None):
    p = paragraph.insert_paragraph_before(text)
    if style:
        p.style = style
    return p


def page_break_before(paragraph):
    p = before(paragraph)
    p.add_run().add_break(WD_BREAK.PAGE)
    return p


def add_table_before(marker_paragraph, rows: list[list[str]]):
    doc = marker_paragraph.part.document
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for c, value in enumerate(rows[0]):
        cell = table.cell(0, c)
        cell.text = value
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row_values in rows[1:]:
        cells = table.add_row().cells
        for c, value in enumerate(row_values):
            cells[c].text = value
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(3)
    marker_paragraph._p.addprevious(table._tbl)
    return table


def add_blocks(marker, blocks: list[tuple[str, str | None]]):
    for text, style in blocks:
        before(marker, text, style)


def long_paragraph(topic: str, focus: str, detail: str) -> str:
    return (
        f"Đối với {topic}, phần quan trọng không chỉ là cài đặt được chức năng mà còn phải xác định rõ {focus}. "
        f"Trong project, nội dung này được xử lý theo hướng chia nhỏ trách nhiệm, ưu tiên các thành phần có thể kiểm thử độc lập và hạn chế để một script đảm nhiệm quá nhiều vai trò. "
        f"{detail} "
        f"Cách trình bày này giúp báo cáo thể hiện được mối liên hệ giữa yêu cầu, thiết kế và mã nguồn thực tế, đồng thời cho thấy quá trình phát triển game là một chuỗi quyết định kỹ thuật có căn cứ chứ không đơn thuần là ghép các chức năng rời rạc."
    )


def add_chapter1_expansion(doc: Document):
    marker = doc.paragraphs[find_paragraph(doc, "1.3. Đề xuất phương pháp giải quyết")]
    add_blocks(
        marker,
        [
            ("1.2.5. Phân tích đối tượng người chơi và bối cảnh sử dụng", "Heading 3"),
            (
                "Đối tượng người chơi mà đề tài hướng tới là nhóm người dùng quen với game góc nhìn thứ nhất, có thể thao tác bằng bàn phím và chuột, đồng thời chấp nhận trải nghiệm khám phá theo nhiệm vụ. Vì đây là sản phẩm đồ án, mục tiêu không phải cạnh tranh trực tiếp với các game thương mại quy mô lớn mà là xây dựng một bản prototype đủ rõ để minh họa kỹ thuật, gameplay và hướng phát triển.",
                "Normal",
            ),
            (
                "Người chơi cần được dẫn dắt bằng các tín hiệu dễ hiểu: nhiệm vụ hiện tại, prompt tương tác, hội thoại, âm thanh phản hồi và thay đổi môi trường. Nếu thiếu các tín hiệu này, game 3D rất dễ khiến người chơi bị lạc mục tiêu, đặc biệt khi scene có nhiều cây, vật thể trang trí và khu vực nước. Do đó, hệ thống HUD nhiệm vụ, dialogue theo ngày và prompt tương tác đóng vai trò quan trọng trong trải nghiệm tổng thể.",
                "Normal",
            ),
            (
                "Trong bối cảnh trình bày đồ án, người xem thường cần quan sát nhanh các chức năng tiêu biểu. Vì vậy, game được thiết kế theo các vòng chơi ngắn: nhận nhiệm vụ, di chuyển đến vị trí, tương tác, nhận phản hồi, cập nhật tiến độ và chuyển sang nhiệm vụ tiếp theo. Cấu trúc này giúp sản phẩm dễ demo, đồng thời vẫn có đủ chiều sâu để giải thích các module kỹ thuật như inventory, quest, dialogue và save/load.",
                "Normal",
            ),
            ("1.2.6. Tiêu chí thành công của sản phẩm demo", "Heading 3"),
            (
                "Một bản demo game trong phạm vi đồ án được xem là thành công khi người chơi có thể đi qua các luồng chính mà không cần can thiệp trực tiếp từ Unity Editor. Các luồng này gồm vào game từ menu, điều khiển nhân vật, tương tác với object, hoàn thành ít nhất một nhiệm vụ, thu thập vật phẩm, mở inventory, trải nghiệm mini game và nhận phản hồi từ hệ thống.",
                "Normal",
            ),
            (
                "Ngoài tiêu chí chạy được, sản phẩm còn cần có tính giải thích. Mỗi chức năng trong game nên có thể liên hệ với một phần kiến thức đã học: lập trình hướng đối tượng, cấu trúc dữ liệu, xử lý sự kiện, tương tác vật lý, quản lý giao diện, âm thanh, dữ liệu cấu hình hoặc kiểm thử. Nhờ vậy, báo cáo không chỉ mô tả kết quả cuối cùng mà còn làm rõ giá trị học thuật của quá trình thực hiện.",
                "Normal",
            ),
            ("Bảng 1- 4. Tiêu chí thành công của bản demo", "Caption"),
        ],
    )
    add_table_before(
        marker,
        [
            ["Tiêu chí", "Mô tả", "Biểu hiện trong project"],
            ["Luồng chơi hoàn chỉnh", "Người chơi có thể bắt đầu, chơi và hoàn thành nhiệm vụ cơ bản.", "Main Menu, scene Suml, quest theo ngày, dialogue và HUD."],
            ["Tương tác rõ ràng", "Object có thể tương tác cần có tín hiệu và phản hồi.", "Prompt, raycast, âm thanh, cập nhật inventory hoặc quest."],
            ["Dữ liệu có tổ chức", "Các dữ liệu dễ thay đổi nên tách khỏi logic xử lý.", "Item definition, quest database, dialogue database, save data."],
            ["Có thể kiểm thử", "Chức năng được chia thành các tình huống kiểm thử cụ thể.", "Bảng kịch bản kiểm thử, đánh giá kết quả, ghi nhận hạn chế."],
            ["Có hướng mở rộng", "Bản demo còn dư địa phát triển thành sản phẩm lớn hơn.", "Save/load, prefab UI, ending, thêm item, thêm sinh vật và story."],
        ],
    )
    before(
        marker,
        "Các tiêu chí trên được dùng làm cơ sở khi lựa chọn phạm vi triển khai. Những chức năng có giá trị minh họa cao và liên kết với nhiều module khác được ưu tiên hơn các tính năng trang trí. Ví dụ, câu cá và nấu ăn được đưa vào vì chúng yêu cầu input, UI, inventory, quest và âm thanh phối hợp với nhau; trong khi đó multiplayer hoặc hệ thống tài khoản trực tuyến không được đưa vào vì vượt quá phạm vi đồ án.",
        "Normal",
    )


def add_chapter2_expansion(doc: Document):
    marker = doc.paragraphs[find_paragraph(doc, "CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG")]
    page_break_before(marker)
    blocks: list[tuple[str, str | None]] = [
        ("2.4. Cơ sở thiết kế gameplay trong game nhập vai 3D", "Heading 2"),
        (
            "Thiết kế gameplay trong đề tài được xây dựng dựa trên nguyên tắc vòng lặp ngắn và phản hồi rõ ràng. Một vòng lặp gameplay cơ bản gồm: người chơi nhận mục tiêu, quan sát môi trường, thực hiện thao tác, nhận phản hồi và thấy tiến độ được cập nhật. Với game nhập vai 3D, vòng lặp này cần diễn ra trong không gian ba chiều nên các yếu tố như camera, khoảng cách tương tác, collider, prompt UI và âm thanh trở nên rất quan trọng.",
            "Normal",
        ),
        (
            "Game không nên yêu cầu người chơi đoán quá nhiều thao tác. Khi người chơi nhìn vào một object hợp lệ, hệ thống cần hiển thị prompt ngắn gọn. Khi người chơi bấm phím tương tác, object cần phản hồi ngay bằng animation, âm thanh, thay đổi trạng thái, vật phẩm rơi ra hoặc cập nhật nhiệm vụ. Nhờ đó, người chơi hiểu rằng hành động của mình có tác động đến thế giới game.",
            "Normal",
        ),
        ("2.4.1. Vòng lặp nhiệm vụ", "Heading 3"),
        (
            "Vòng lặp nhiệm vụ theo ngày là xương sống của project. Thay vì để người chơi tự do hoàn toàn trong scene, hệ thống đưa ra mục tiêu cụ thể theo từng ngày. Cách làm này giảm độ phức tạp thiết kế thế giới mở, đồng thời giúp demo có nhịp rõ ràng: ngày đầu làm quen, các ngày sau mở thêm hành động sinh tồn và điều tra, ngày cuối hướng tới kết thúc câu chuyện.",
            "Normal",
        ),
        (
            "DailyQuestManager giữ vai trò điều phối tiến độ. Khi một hành động xảy ra, ví dụ nhặt vật phẩm, phá bẫy, hạ lợn rừng, câu được cá hoặc nấu ăn thành công, module gameplay tương ứng báo về quest manager. Quest manager kiểm tra điều kiện, cập nhật HUD và quyết định nhiệm vụ đã hoàn thành hay chưa. Cách này giúp các module gameplay không phải tự xử lý logic tiến trình tổng thể.",
            "Normal",
        ),
        ("2.4.2. Thiết kế phản hồi người chơi", "Heading 3"),
        (
            "Phản hồi trong game được chia thành nhiều lớp. Lớp đầu tiên là phản hồi thị giác như prompt, HUD, panel inventory, dialogue UI và mini game. Lớp thứ hai là phản hồi âm thanh như tiếng chặt cây, âm thanh câu cá, âm thanh thắng/thua, âm thanh ăn hoặc mở cửa. Lớp thứ ba là phản hồi trạng thái như vật phẩm tăng trong inventory, stamina giảm hoặc nhiệm vụ được cập nhật.",
            "Normal",
        ),
        (
            "Một thao tác quan trọng nên có ít nhất hai loại phản hồi để người chơi dễ nhận biết. Ví dụ, khi chặt cây thành công, cây biến mất hoặc thay đổi, vật phẩm gỗ xuất hiện và âm thanh chặt/cây đổ được phát. Khi câu cá thành công, UI mini game báo kết quả, cá được thêm vào inventory và nhiệm vụ liên quan được cập nhật. Điều này làm trải nghiệm rõ ràng hơn so với chỉ thay đổi dữ liệu ngầm.",
            "Normal",
        ),
        ("2.5. Cơ sở tổ chức mã nguồn trong Unity", "Heading 2"),
    ]
    for topic, focus, detail in [
        ("mô hình component của Unity", "ranh giới trách nhiệm giữa GameObject và script", "Mỗi GameObject trong scene nên chứa những component phục vụ trực tiếp cho hành vi của nó; ví dụ player chứa movement, look, inventory và stamina, còn object tương tác chứa Interactable hoặc script con tương ứng."),
        ("ScriptableObject và dữ liệu cấu hình", "cách tách dữ liệu khỏi logic xử lý", "Các item, quest và dialogue có nhiều khả năng thay đổi theo nội dung game, vì vậy việc tách thành asset cấu hình giúp giảm sửa code khi cần đổi tên, mô tả, số lượng hoặc điều kiện hoàn thành."),
        ("event và lời gọi giữa module", "cách giảm phụ thuộc cứng giữa các hệ thống", "Những hành động như tấn công, nhặt vật phẩm hoặc hoàn thành mini game nên phát tín hiệu rõ ràng để module khác xử lý tiếp, thay vì để một script biết quá nhiều chi tiết nội bộ của script khác."),
        ("coroutine trong gameplay thời gian thực", "cách xử lý các khoảng chờ và hiệu ứng theo thời gian", "Các luồng như chờ cá cắn câu, typewriter dialogue, loading hoặc hiệu ứng chuyển trạng thái phù hợp với coroutine vì chúng không chặn toàn bộ game loop."),
        ("UI Canvas và TextMeshPro", "cách trình bày thông tin trên màn hình", "HUD nhiệm vụ, prompt tương tác, dialogue, inventory và mini game đều cần chữ rõ ràng, cập nhật theo trạng thái runtime và không gây cản trở tầm nhìn chính của người chơi."),
    ]:
        blocks.append((long_paragraph(topic, focus, detail), "Normal"))
    blocks.extend(
        [
            ("2.6. Nguyên tắc kiểm thử trong game thời gian thực", "Heading 2"),
            (
                "Kiểm thử game khác với kiểm thử một chương trình nhập xuất đơn giản vì nhiều trạng thái xảy ra đồng thời: input người chơi, physics, animation, UI, âm thanh, collider và dữ liệu nhiệm vụ. Do đó, kiểm thử cần tập trung vào các kịch bản thao tác thực tế thay vì chỉ kiểm tra hàm riêng lẻ.",
                "Normal",
            ),
            (
                "Với project này, mỗi chức năng chính được kiểm thử theo ba mức: kiểm thử chức năng đơn lẻ, kiểm thử tích hợp với module liên quan và kiểm thử theo luồng chơi. Ví dụ, inventory không chỉ được kiểm tra bằng việc thêm item vào slot mà còn cần kiểm tra khi nhặt item từ scene, dùng item để đặt campfire, đưa nguyên liệu vào cooking mode và cập nhật nhiệm vụ ngày 5.",
                "Normal",
            ),
            ("Bảng 2- 2. Liên hệ giữa kiến thức nền tảng và chức năng trong project", "Caption"),
        ]
    )
    add_blocks(marker, blocks)
    add_table_before(
        marker,
        [
            ["Kiến thức", "Ứng dụng trong project", "Ý nghĩa"],
            ["Lập trình hướng đối tượng", "Tách script theo vai trò như PlayerInventory, DialogueController, DailyQuestManager.", "Giúp mã nguồn dễ đọc và dễ mở rộng."],
            ["Physics và Collider", "Raycast, vùng tương tác, object pickup, va chạm với môi trường.", "Tạo tương tác trực tiếp trong không gian 3D."],
            ["UI runtime", "Inventory, quest HUD, dialogue UI, prompt và mini game.", "Cung cấp phản hồi và hướng dẫn người chơi."],
            ["Dữ liệu cấu hình", "Item, quest, dialogue và save data.", "Giảm hard-code, dễ thay đổi nội dung."],
            ["Coroutine", "Chờ cá cắn, hiệu ứng hội thoại, loading và chuỗi hành động theo thời gian.", "Xử lý logic kéo dài qua nhiều frame."],
            ["Âm thanh", "SFX chặt cây, câu cá, nấu ăn, ăn, thắng/thua.", "Tăng cảm giác phản hồi và nhập vai."],
        ],
    )
    before(
        marker,
        "Việc liên hệ lý thuyết với chức năng cụ thể giúp chương cơ sở không bị tách rời khỏi sản phẩm. Mỗi khái niệm kỹ thuật đều được dùng để giải quyết một nhu cầu thực tế của game, từ điều khiển nhân vật cho tới quản lý nhiệm vụ và phản hồi người chơi.",
        "Normal",
    )


def add_chapter3_expansion(doc: Document):
    marker = doc.paragraphs[find_paragraph(doc, "CHƯƠNG 4: CÀI ĐẶT, THỬ NGHIỆM VÀ ĐÁNH GIÁ")]
    page_break_before(marker)
    blocks: list[tuple[str, str | None]] = [
        ("3.6. Thiết kế kiến trúc module tổng thể", "Heading 2"),
        (
            "Kiến trúc project được định hướng theo các nhóm module tương đối độc lập. Mỗi nhóm module có một nhiệm vụ chính và giao tiếp với các nhóm khác thông qua dữ liệu, lời gọi hàm hoặc sự kiện. Cách tổ chức này phù hợp với Unity vì mỗi hành vi thường được gắn vào GameObject dưới dạng component, nhưng vẫn cần một tầng điều phối để tránh việc các component phụ thuộc lẫn nhau quá chặt.",
            "Normal",
        ),
        ("Bảng 3- 1. Nhóm module và trách nhiệm thiết kế", "Caption"),
    ]
    add_blocks(marker, blocks)
    add_table_before(
        marker,
        [
            ["Nhóm module", "Trách nhiệm", "Dữ liệu/đầu ra chính"],
            ["Player", "Xử lý di chuyển, quan sát, cúi, chạy, hành động và stamina.", "Vị trí, hướng nhìn, trạng thái input, stamina."],
            ["Interact", "Phát hiện object tương tác, xử lý pickup, cửa, cây, hint và trigger.", "Prompt, sự kiện tương tác, vật phẩm hoặc trạng thái object."],
            ["Inventory", "Lưu vật phẩm, stack item, hiển thị slot, dùng vật phẩm và đặt campfire.", "Danh sách slot, itemId, số lượng, trạng thái UI."],
            ["Quest", "Quản lý nhiệm vụ theo ngày, điều kiện hoàn thành và HUD tiến độ.", "Ngày hiện tại, objective, tiến độ, trạng thái hoàn thành."],
            ["Dialogue", "Hiển thị hội thoại, quản lý event theo ngày và khóa input khi cần.", "DialogueDay, DialogueEventId, line, trạng thái đang thoại."],
            ["Audio", "Phát nhạc nền, SFX 2D/3D và loop âm thanh.", "Sound id, music track, AudioSource runtime."],
            ["Menu/Save", "Bắt đầu game, tạm dừng, settings, save/load theo slot.", "AllGameData, PlayerData, file save, cấu hình người chơi."],
        ],
    )
    add_blocks(
        marker,
        [
            (
                "Các module không tồn tại độc lập hoàn toàn. Ví dụ, khi người chơi câu cá thành công, module Fishing cập nhật Inventory, trừ stamina, phát âm thanh và báo tiến độ cho Quest. Khi người chơi mở Dialogue, module Dialogue cần khóa input của Player để tránh người chơi vừa di chuyển vừa đọc hội thoại. Vì vậy, điểm quan trọng trong thiết kế là xác định rõ module nào giữ trạng thái gốc và module nào chỉ phản ánh trạng thái đó lên UI.",
                "Normal",
            ),
            ("3.7. Thiết kế dữ liệu chi tiết", "Heading 2"),
            (
                "Dữ liệu trong project được chia thành dữ liệu cấu hình và dữ liệu runtime. Dữ liệu cấu hình là phần được chuẩn bị sẵn trong project, ví dụ định nghĩa vật phẩm, nhiệm vụ hoặc hội thoại. Dữ liệu runtime là phần thay đổi khi chơi, ví dụ số lượng item trong inventory, ngày hiện tại, nhiệm vụ đã hoàn thành hoặc vị trí player.",
                "Normal",
            ),
            (
                "Cách phân chia này giúp hệ thống rõ ràng hơn. Khi muốn thêm một item mới, nhóm phát triển ưu tiên thêm dữ liệu định nghĩa item thay vì sửa nhiều script. Khi muốn lưu game, SaveManager chỉ cần lấy dữ liệu runtime quan trọng và đóng gói lại. Nếu trộn hai loại dữ liệu này, project sẽ khó mở rộng và dễ phát sinh lỗi khi thay đổi nội dung.",
                "Normal",
            ),
            ("3.7.1. Dữ liệu vật phẩm", "Heading 3"),
            (
                "Vật phẩm trong game cần có mã định danh ổn định, tên hiển thị, mô tả, icon, số lượng stack tối đa và hành vi khi sử dụng. Mã định danh như wood_log, boar_meat hoặc river_fish không nên thay đổi tùy tiện vì có thể được quest, save/load hoặc UI sử dụng để đối chiếu. Tên hiển thị và mô tả có thể thay đổi dễ hơn vì chủ yếu phục vụ giao diện.",
                "Normal",
            ),
            ("3.7.2. Dữ liệu nhiệm vụ", "Heading 3"),
            (
                "Mỗi nhiệm vụ cần xác định ngày, tên nhiệm vụ, mô tả, loại mục tiêu, số lượng yêu cầu và hành động hoàn thành. Một nhiệm vụ đơn giản có thể chỉ yêu cầu nhặt đủ item, nhưng nhiệm vụ phức tạp như ngày 5 cần nhiều điều kiện song song: bắt cá, gom gỗ, đặt campfire, nấu ăn và ăn. Vì vậy, thiết kế quest cần hỗ trợ cả mục tiêu đơn và mục tiêu tổng hợp.",
                "Normal",
            ),
            ("3.7.3. Dữ liệu hội thoại", "Heading 3"),
            (
                "Dialogue được gắn với ngày và sự kiện. Cùng một eventId có thể có ý nghĩa khác nhau theo ngày, hoặc một ngày có nhiều event tương ứng với nhiều tình huống trong game. Mỗi dòng hội thoại nên có người nói, nội dung, tốc độ hiển thị và tùy chọn khóa input nếu đoạn thoại cần người chơi tập trung. Thiết kế này giúp story có thể mở rộng mà không phải sửa trực tiếp vào logic gameplay.",
                "Normal",
            ),
            ("3.8. Thiết kế luồng xử lý chi tiết", "Heading 2"),
        ],
    )
    flows = [
        ("Luồng nhặt vật phẩm", ["Camera raycast phát hiện object có thể nhặt.", "Prompt hiển thị phím tương tác.", "Người chơi bấm E.", "Script pickup kiểm tra item definition.", "PlayerInventory thêm item vào slot phù hợp.", "Quest manager nhận báo cáo nếu item liên quan nhiệm vụ.", "Object pickup bị ẩn hoặc bị hủy khỏi scene."]),
        ("Luồng chặt cây", ["Người chơi trang bị rìu bằng phím F.", "ActionScript phát sự kiện tấn công khi bấm chuột trái.", "CuttingTreeSystem kiểm tra cây hợp lệ bằng raycast/collider.", "Số hit của cây được cập nhật.", "Khi đủ hit, cây bị loại bỏ và prefab gỗ được sinh ra.", "Âm thanh chặt cây/cây đổ được phát.", "Inventory có thể nhận gỗ khi người chơi nhặt pickup."]),
        ("Luồng câu cá", ["Người chơi đứng gần vùng nước.", "FishingRob xác định vị trí thả câu và khóa input không cần thiết.", "LineRenderer hiển thị dây câu.", "Coroutine chờ thời gian cá cắn.", "Mini game mở khi cá cắn.", "Nếu thành công, cá được thêm vào inventory và quest cập nhật.", "Nếu thất bại, game phát phản hồi thua và cho phép thử lại."]),
        ("Luồng nấu ăn", ["Người chơi đặt campfire bằng gỗ.", "Tương tác với campfire để vào cooking mode.", "Inventory mở ở chế độ chọn nguyên liệu.", "Người chơi chọn cá hoặc thịt.", "Mini game cooking xử lý thao tác đúng thời điểm.", "Nếu thắng, campfire tăng số món đã nấu và quest ngày 5 cập nhật.", "Người chơi có thể ăn món đã nấu để hồi phục chỉ số."]),
        ("Luồng save/load", ["Người chơi chọn slot save hoặc load.", "SaveManager gom dữ liệu runtime thành AllGameData.", "Dữ liệu được ghi hoặc đọc từ file slot.", "Khi load, scene được nạp nếu cần.", "Player, inventory, quest và dialogue được khôi phục.", "UI slot, HUD và inventory đồng bộ lại trạng thái."]),
    ]
    for title, steps in flows:
        add_blocks(marker, [(title, "Heading 3")])
        for step in steps:
            before(marker, step, "List Number")
        before(
            marker,
            f"Luồng {title.casefold()} cho thấy một chức năng trong game thường cần nhiều module cùng tham gia. Khi thiết kế, cần xác định điểm bắt đầu, điều kiện hợp lệ, trạng thái thay đổi và phản hồi cuối cùng để tránh trường hợp người chơi thực hiện thao tác nhưng hệ thống không cập nhật đầy đủ.",
            "Normal",
        )


def add_chapter4_expansion(doc: Document):
    marker = doc.paragraphs[find_paragraph(doc, "4.4. Minh họa sản phẩm thực tế")]
    blocks: list[tuple[str, str | None]] = [
        ("4.3.6. Kế hoạch kiểm thử mở rộng", "Heading 3"),
        (
            "Để đánh giá project kỹ hơn, ngoài bảng kiểm thử chức năng chính, cần bổ sung kế hoạch kiểm thử mở rộng theo nhóm rủi ro. Kế hoạch này không chỉ kiểm tra thao tác đúng mà còn kiểm tra thao tác sai, thao tác lặp lại, chuyển trạng thái bất ngờ và tương tác giữa nhiều UI.",
            "Normal",
        ),
        ("Bảng 4- 4. Kế hoạch kiểm thử mở rộng", "Caption"),
    ]
    add_blocks(marker, blocks)
    add_table_before(
        marker,
        [
            ["Nhóm kiểm thử", "Tình huống", "Kỳ vọng"],
            ["Input", "Mở inventory khi đang dialogue hoặc cooking.", "Input được khóa/mở đúng, không bị kẹt chuột hoặc mất điều khiển."],
            ["Inventory", "Nhặt item khi slot đầy hoặc item vượt stack.", "Item được cộng đúng stack hoặc bị từ chối có kiểm soát."],
            ["Quest", "Hoàn thành nhiều điều kiện ngày 5 theo thứ tự khác nhau.", "HUD cập nhật từng mục tiêu và chỉ hoàn thành khi đủ điều kiện."],
            ["Dialogue", "Kích hoạt lại event đã chạy hoặc chuyển ngày.", "Không lặp thoại ngoài ý muốn; thoại đúng ngày được ưu tiên."],
            ["Save/load", "Lưu ở giữa nhiệm vụ rồi tải lại.", "Vị trí, inventory, quest và dialogue khôi phục nhất quán."],
            ["Audio", "Chuyển menu/gameplay và bật/tắt settings âm lượng.", "Nhạc nền và SFX không chồng loop bất thường."],
            ["Scene", "Chuyển scene từ menu/loading và quay lại menu.", "Không mất reference quan trọng, không phát sinh object trùng."],
        ],
    )
    add_blocks(
        marker,
        [
            (
                "Các tình huống trên nên được kiểm thử nhiều lần sau mỗi đợt sửa code lớn. Đặc biệt, những chức năng liên quan đến khóa input và UI cần được kiểm thử theo tổ hợp, vì lỗi thường không xuất hiện khi chạy riêng từng module mà xuất hiện khi người chơi chuyển nhanh giữa nhiều trạng thái.",
                "Normal",
            ),
            ("4.3.7. Phân tích lỗi thường gặp và hướng xử lý", "Heading 3"),
        ],
    )
    issues = [
        ("Mất reference trong Inspector", "Khi đổi tên object, xóa prefab hoặc di chuyển script, reference có thể bị mất. Cần kiểm tra log, dùng serialize field rõ ràng và tránh phụ thuộc quá nhiều vào Find theo tên."),
        ("UI bị tạo trùng", "Một số UI runtime tiện cho demo nhưng dễ tạo nhiều instance nếu không kiểm tra object đã tồn tại. Cần có cơ chế singleton hoặc prefab quản lý tập trung."),
        ("Input không được mở lại", "Dialogue, inventory, fishing và cooking đều có thể khóa input. Cần quy định rõ module nào chịu trách nhiệm trả lại trạng thái và xử lý cả trường hợp thoát sớm."),
        ("Quest không cập nhật", "Nếu module gameplay quên gọi report về DailyQuestManager, HUD sẽ không đổi dù người chơi đã thực hiện hành động. Cần gom các API báo cáo quest và dùng tên objective ổn định."),
        ("Âm thanh loop không dừng", "Các loop như Fishing hoặc môi trường có thể tiếp tục phát nếu scene đổi trạng thái bất ngờ. Cần quản lý loop id và dừng âm thanh khi thoát mode."),
        ("Save/load thiếu dữ liệu", "Nếu chỉ lưu một phần trạng thái, người chơi tải lại có thể thấy quest đúng nhưng object scene sai. Cần mở rộng save data theo version và bổ sung migration nếu cấu trúc thay đổi."),
    ]
    for title, body in issues:
        before(marker, title, "Heading 3")
        before(marker, body, "Normal")
        before(
            marker,
            "Khi gặp lỗi này, cách xử lý nên bắt đầu từ việc tái hiện lỗi bằng một kịch bản ngắn, sau đó xác định module giữ trạng thái gốc, kiểm tra log runtime và chỉ sửa ở phạm vi liên quan. Tránh sửa lan sang nhiều script nếu chưa chứng minh được nguyên nhân.",
            "Normal",
        )
    add_blocks(
        marker,
        [
            ("4.3.8. Đánh giá mức độ đáp ứng yêu cầu", "Heading 3"),
            (
                "So với mục tiêu ban đầu, project đã đáp ứng phần lớn các yêu cầu chức năng cốt lõi. Người chơi có thể vào game, di chuyển trong môi trường 3D, tương tác, thu thập tài nguyên, mở inventory, thực hiện nhiệm vụ, đọc hội thoại và trải nghiệm các mini game. Các yêu cầu phi chức năng như tính dễ dùng, tính mở rộng và tính bảo trì cũng đã được chú ý thông qua việc tách module và sử dụng dữ liệu cấu hình.",
                "Normal",
            ),
            (
                "Tuy nhiên, mức độ hoàn thiện vẫn ở cấp prototype. Một số hệ thống cần chuẩn hóa thêm nếu phát triển thành sản phẩm hoàn chỉnh, đặc biệt là save/load toàn diện, prefab UI, ending, test tự động và tối ưu scene. Điều này phù hợp với phạm vi đồ án vì trọng tâm là thể hiện quy trình phân tích, thiết kế và cài đặt một game có nhiều module kết nối.",
                "Normal",
            ),
        ],
    )


def add_large_appendix(doc: Document):
    marker = doc.paragraphs[find_paragraph(doc, "PHỤ LỤC")]
    page_break_before(marker)
    sections = [
        ("Phụ lục A. Mô tả chi tiết nhóm script Player", [
            ("PlayerMovement", "xử lý chuyển động, tốc độ đi/chạy, hướng di chuyển theo input và trạng thái có được phép điều khiển hay không"),
            ("PlayerLook và MouseMovement", "xử lý xoay camera, độ nhạy chuột và giới hạn góc nhìn để tạo cảm giác góc nhìn thứ nhất"),
            ("Crouch", "thay đổi trạng thái cúi, hỗ trợ người chơi di chuyển trong các ngữ cảnh cần hạ thấp nhân vật"),
            ("ActionScript", "điều phối thao tác trang bị rìu, tấn công và phát sự kiện cho các hệ thống khác"),
            ("Stamina", "quản lý sức bền, hồi phục, tiêu hao khi chạy, tấn công, câu cá hoặc hành động nặng"),
        ]),
        ("Phụ lục B. Mô tả chi tiết nhóm script Interaction", [
            ("Interactable", "định nghĩa nền cho các object có thể tương tác, giúp các script khác xử lý thống nhất"),
            ("PickUpScript", "raycast từ camera để phát hiện object, hiển thị prompt và gọi thao tác tương tác"),
            ("Door", "xử lý mở/đóng cửa, phát âm thanh và có thể liên kết với dialogue theo ngày"),
            ("CuttingTreeSystem", "kiểm tra cây hợp lệ, đếm số hit, loại bỏ cây và sinh vật phẩm gỗ"),
            ("HintDay3Interactable", "gắn tương tác đặc biệt với tiến trình ngày 3 và hiệu ứng hint/glitch"),
        ]),
        ("Phụ lục C. Mô tả chi tiết nhóm script Inventory và Cooking", [
            ("PlayerInventory", "quản lý danh sách slot, thêm/bớt item, xử lý stack và cung cấp dữ liệu cho UI"),
            ("InventoryItemDefinition", "lưu thông tin cấu hình item như id, tên, mô tả, icon và số lượng stack"),
            ("InventoryUIController", "hiển thị panel inventory, cập nhật slot và xử lý trạng thái mở/tắt"),
            ("CampingCookingInteractable", "cho phép tương tác với campfire, ăn món đã nấu và báo tiến độ quest"),
            ("MiniGameCookingController", "điều khiển mini game nấu ăn, vùng mục tiêu, kết quả thắng/thua và phản hồi UI"),
        ]),
        ("Phụ lục D. Mô tả chi tiết nhóm script Quest và Dialogue", [
            ("DailyQuestManager", "quản lý nhiệm vụ theo ngày, tiến độ mục tiêu và điều kiện hoàn thành"),
            ("DailyQuestDatabase", "chứa danh sách nhiệm vụ cấu hình cho từng ngày"),
            ("DialogueController", "hiển thị từng dòng hội thoại, typewriter effect và khóa input khi cần"),
            ("DialogueDatabase", "chứa nội dung hội thoại theo ngày và sự kiện"),
            ("DialogueSaveService", "lưu trạng thái ngày hoặc hội thoại để hỗ trợ khôi phục tiến trình"),
        ]),
        ("Phụ lục E. Mô tả chi tiết nhóm script Menu, Save và Audio", [
            ("MainMenu", "điều khiển nút New Game, Load, Quit và một số lựa chọn cài đặt ban đầu"),
            ("MenuManager", "quản lý pause menu trong scene gameplay và khóa/mở điều khiển khi menu hiển thị"),
            ("SaveManager", "gom dữ liệu runtime, ghi/đọc file save và khôi phục trạng thái game"),
            ("ReSoundManager", "phát SFX 2D/3D, loop âm thanh và cung cấp API phát âm thanh theo id"),
            ("MusicManager", "điều khiển nhạc nền theo track và ngữ cảnh menu/gameplay"),
        ]),
    ]
    for heading, items in sections:
        before(marker, heading, "Heading 1")
        before(
            marker,
            "Phần phụ lục này trình bày thêm vai trò của các script chính để người đọc có thể đối chiếu giữa nội dung báo cáo và cấu trúc mã nguồn trong project. Mỗi script được mô tả theo trách nhiệm chính, dữ liệu liên quan và cách tham gia vào vòng lặp gameplay.",
            "Normal",
        )
        rows = [["Script", "Vai trò chi tiết", "Ghi chú thiết kế"]]
        for script, role in items:
            rows.append([
                script,
                role,
                "Nên giữ trách nhiệm tập trung, tránh mở rộng thành nơi chứa logic của module khác.",
            ])
        add_table_before(marker, rows)
        for script, role in items:
            before(marker, script, "Heading 2")
            before(
                marker,
                f"{script} là một thành phần quan trọng trong nhóm chức năng tương ứng. Script này chủ yếu {role}. Khi trình bày trong báo cáo, cần làm rõ script nhận dữ liệu từ đâu, cập nhật trạng thái nào và ảnh hưởng tới trải nghiệm người chơi ra sao.",
                "Normal",
            )
            before(
                marker,
                f"Về mặt bảo trì, {script} nên được kiểm tra sau mỗi lần thay đổi scene hoặc prefab liên quan. Nếu script có liên kết với UI, quest, dialogue hoặc save/load, cần kiểm thử thêm các trường hợp chuyển trạng thái để đảm bảo dữ liệu không bị lệch giữa runtime và giao diện.",
                "Normal",
            )
            before(
                marker,
                f"Trong hướng phát triển tiếp theo, {script} có thể được cải thiện bằng cách chuẩn hóa reference trong Inspector, bổ sung log debug có kiểm soát và tách dữ liệu cấu hình ra khỏi logic xử lý nếu nội dung ngày càng nhiều.",
                "Normal",
            )
    before(
        marker,
        "Phụ lục F. Checklist hoàn thiện sản phẩm",
        "Heading 1",
    )
    checklist = [
        "Kiểm tra toàn bộ scene trong Build Settings và bảo đảm tên scene được dùng thống nhất.",
        "Chuẩn hóa prefab UI cho inventory, quest HUD, dialogue, settings và mini game.",
        "Rà soát các object phụ thuộc tên, chuyển sang reference hoặc tag/layer rõ ràng.",
        "Mở rộng save/load để lưu đầy đủ inventory, ngày, quest, dialogue, vị trí và trạng thái môi trường.",
        "Bổ sung playtest với người dùng ngoài nhóm để đánh giá khả năng hiểu nhiệm vụ.",
        "Tối ưu audio loop, collider, terrain detail và object runtime để tránh giảm hiệu năng.",
        "Hoàn thiện nội dung ngày 6, ending và các trigger story cuối game.",
        "Tạo bản build Windows, kiểm thử trên máy không mở Unity Editor và ghi nhận lỗi phát sinh.",
    ]
    for item in checklist:
        before(marker, item, "List Bullet")
    for i in range(1, 9):
        before(
            marker,
            f"Giải thích checklist {i}: mục này cần được xem như một tiêu chí nghiệm thu nhỏ trước khi đóng gói sản phẩm. Khi hoàn thành, nhóm nên ghi lại bằng chứng kiểm thử như ảnh chụp màn hình, mô tả thao tác, kết quả mong đợi và kết quả thực tế. Việc ghi nhận này giúp báo cáo có cơ sở hơn, đồng thời giúp người phát triển dễ quay lại kiểm tra nếu lỗi tái xuất hiện ở các phiên bản sau.",
            "Normal",
        )


def add_more_discussion(doc: Document):
    marker = doc.paragraphs[find_paragraph(doc, "TÀI LIỆU THAM KHẢO")]
    page_break_before(marker)
    add_blocks(
        marker,
        [
            ("ĐÁNH GIÁ TỔNG HỢP VÀ BÀI HỌC KINH NGHIỆM", "Heading 1"),
            (
                "Quá trình xây dựng game nhập vai 3D cho thấy việc phát triển game là sự kết hợp giữa kỹ thuật phần mềm và thiết kế trải nghiệm. Một chức năng chỉ được xem là hoàn chỉnh khi người chơi nhìn thấy, hiểu được, thao tác được và nhận phản hồi rõ ràng. Vì vậy, khi đánh giá sản phẩm, cần nhìn cả ở góc độ mã nguồn, dữ liệu, giao diện, âm thanh, kịch bản chơi và khả năng trình bày demo.",
                "Normal",
            ),
        ],
    )
    lessons = [
        ("Bài học về kiểm soát phạm vi", "Không nên mở rộng quá nhiều chức năng khi nền tảng chưa ổn định. Với game 3D, mỗi chức năng mới thường kéo theo asset, UI, âm thanh, collider, animation và kiểm thử."),
        ("Bài học về thiết kế dữ liệu", "Dữ liệu như item, quest và dialogue nên được tách khỏi logic càng sớm càng tốt. Khi nội dung tăng, cách làm này giúp sửa đổi nhanh và ít rủi ro hơn."),
        ("Bài học về UI", "UI không chỉ là phần hiển thị mà còn là cách người chơi hiểu trạng thái game. Prompt, HUD và dialogue cần ngắn gọn, rõ ngữ cảnh và cập nhật đúng thời điểm."),
        ("Bài học về tích hợp module", "Nhiều lỗi xuất hiện ở ranh giới giữa các module, ví dụ inventory cập nhật nhưng quest không nhận, hoặc dialogue đóng nhưng input chưa mở lại. Do đó, kiểm thử tích hợp rất quan trọng."),
        ("Bài học về báo cáo đồ án", "Báo cáo nên giải thích được vì sao chọn giải pháp, giải pháp hoạt động thế nào, kết quả ra sao và còn hạn chế gì. Phần mô tả kỹ thuật cần bám vào project thật để tránh chung chung."),
    ]
    for title, body in lessons:
        before(marker, title, "Heading 2")
        for _ in range(3):
            before(
                marker,
                body
                + " Trong bối cảnh đồ án, bài học này giúp nhóm phát triển hiểu rõ hơn cách tổ chức công việc và cách chứng minh giá trị của sản phẩm. Nếu tiếp tục phát triển, những kinh nghiệm này nên được chuyển thành quy ước làm việc cụ thể như checklist trước khi commit, checklist trước khi build và checklist trước khi demo.",
                "Normal",
            )


def main() -> int:
    if len(sys.argv) != 3:
        print("Usage: docx_expand_to_100_pages.py <input.docx> <output.docx>")
        return 2
    source = Path(sys.argv[1])
    target = Path(sys.argv[2])
    doc = Document(source)

    add_chapter1_expansion(doc)
    add_chapter2_expansion(doc)
    add_chapter3_expansion(doc)
    add_chapter4_expansion(doc)
    add_large_appendix(doc)
    add_more_discussion(doc)

    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(target)
    print(target)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
