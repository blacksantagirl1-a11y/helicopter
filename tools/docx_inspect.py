from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")


def clean(text: str) -> str:
    return " ".join((text or "").split())


def main() -> int:
    if len(sys.argv) != 2:
        print("Usage: docx_inspect.py <docx>")
        return 2

    path = Path(sys.argv[1])
    doc = Document(path)

    print(f"FILE: {path}")
    print(f"PARAGRAPHS: {len(doc.paragraphs)}")
    print(f"TABLES: {len(doc.tables)}")
    print()

    print("HEADINGS")
    for i, p in enumerate(doc.paragraphs):
        text = clean(p.text)
        if not text:
            continue
        style = p.style.name if p.style is not None else ""
        if "Heading" in style or "Title" in style or text.isupper():
            print(f"{i:04d} | {style} | {text[:180]}")

    print()
    print("PARAGRAPHS")
    for i, p in enumerate(doc.paragraphs):
        text = clean(p.text)
        if not text:
            continue
        style = p.style.name if p.style is not None else ""
        print(f"{i:04d} | {style} | {text[:260]}")

    print()
    print("TABLES")
    for ti, table in enumerate(doc.tables):
        print(f"TABLE {ti}: rows={len(table.rows)} cols={len(table.columns)}")
        for ri, row in enumerate(table.rows[:8]):
            cells = [clean(c.text)[:90] for c in row.cells]
            print(f"  R{ri}: " + " | ".join(cells))
        if len(table.rows) > 8:
            print("  ...")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
