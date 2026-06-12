from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def normalize(text: str) -> str:
    return " ".join((text or "").split()).casefold()


def find_paragraph(doc: Document, needle: str) -> int:
    target = normalize(needle)
    for i, paragraph in enumerate(doc.paragraphs):
        if normalize(paragraph.text) == target:
            return i
    raise ValueError(f"Paragraph not found: {needle}")


def insert_before(paragraph, text: str = "", style: str | None = None):
    new_paragraph = paragraph.insert_paragraph_before(text)
    if style:
        new_paragraph.style = style
    return new_paragraph


def insert_block_before(doc: Document, marker: str, entries: list[tuple[str, str | None]]):
    marker_paragraph = doc.paragraphs[find_paragraph(doc, marker)]
    inserted = []
    for text, style in entries:
        inserted.append(insert_before(marker_paragraph, text, style))
    return inserted


def paragraph_after(paragraph, text: str, style: str | None = None):
    # python-docx has no public insert-after API; inserting before the following
    # paragraph is handled by the caller for deterministic placement.
    raise NotImplementedError


def set_table_style(table):
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.space_before = Pt(0)


def add_table_before(marker_paragraph, rows: list[list[str]]):
    document = marker_paragraph.part.document
    table = document.add_table(rows=1, cols=len(rows[0]))
    for c, value in enumerate(rows[0]):
        table.cell(0, c).text = value
        for run in table.cell(0, c).paragraphs[0].runs:
            run.bold = True
    for row_values in rows[1:]:
        cells = table.add_row().cells
        for c, value in enumerate(row_values):
            cells[c].text = value
    set_table_style(table)
    marker_paragraph._p.addprevious(table._tbl)
    return table


def add_timeline(doc: Document):
    marker = doc.paragraphs[find_paragraph(doc, "1.2. Khảo sát thực trạng")]
    blocks = [
        ("1.1.4. Nội dung thực hiện", "Heading 3"),
        (
            "Trong quá trình thực hiện đồ án, nhóm chia công việc thành các giai đoạn nhỏ để dễ kiểm soát phạm vi, kiểm thử từng chức năng và hoàn thiện sản phẩm theo hướng có thể chơi thử. Nội dung thực hiện không chỉ bao gồm lập trình gameplay mà còn gồm thiết kế scene, chuẩn bị asset, cấu hình dữ liệu, kiểm thử, chụp hình minh họa và hoàn thiện báo cáo.",
            "Normal",
        ),
        ("Bảng 1- 3. Nội dung thực hiện theo giai đoạn", "Caption"),
    ]
    for text, style in blocks:
        insert_before(marker, text, style)

    rows = [
        ["Giai đoạn", "Nội dung chính", "Kết quả đạt được"],
        [
            "Khảo sát và xác định phạm vi",
            "Lựa chọn hướng game nhập vai 3D kết hợp điều tra và sinh tồn; xác định các chức năng cốt lõi phù hợp với thời gian đồ án.",
            "Có ý tưởng tổng thể, vòng lặp gameplay và danh sách chức năng ưu tiên.",
        ],
        [
            "Xây dựng nền tảng project",
            "Thiết lập Unity, scene chính, terrain, camera, nhân vật, hệ thống input và cấu trúc thư mục script.",
            "Tạo được môi trường chơi cơ bản và khung kỹ thuật ban đầu.",
        ],
        [
            "Cài đặt gameplay cốt lõi",
            "Phát triển điều khiển người chơi, raycast interaction, chặt cây, săn lợn rừng, câu cá, inventory và nấu ăn.",
            "Hình thành vòng lặp thu thập tài nguyên và xử lý thử thách.",
        ],
        [
            "Cài đặt tiến trình nhiệm vụ",
            "Xây dựng hội thoại theo ngày, quest HUD, điều kiện hoàn thành nhiệm vụ, clue vision và thay đổi môi trường theo ngày.",
            "Game có tiến trình chơi rõ ràng, người chơi được định hướng bằng nhiệm vụ và hội thoại.",
        ],
        [
            "Hoàn thiện và đánh giá",
            "Bổ sung âm thanh, menu, settings, loading, save/load ở mức demo; playtest các luồng chính và ghi nhận hạn chế.",
            "Có bản demo có thể trình bày, kèm báo cáo mô tả phân tích, thiết kế, cài đặt và đánh giá.",
        ],
    ]
    add_table_before(marker, rows)

    insert_before(
        marker,
        "Việc chia giai đoạn như trên giúp nhóm kiểm soát rủi ro của đề tài 3D. Những chức năng có độ phụ thuộc cao như inventory, quest, dialogue và save/load được phát triển sau khi các thao tác cơ bản đã ổn định, nhờ đó quá trình tích hợp ít phát sinh lỗi dây chuyền hơn.",
        "Normal",
    )


def add_save_design(doc: Document):
    marker = doc.paragraphs[find_paragraph(doc, "CHƯƠNG 4: CÀI ĐẶT, THỬ NGHIỆM VÀ ĐÁNH GIÁ")]
    for text, style in [
        ("3.5. Thiết kế lưu trữ dữ liệu và trạng thái game", "Heading 2"),
        (
            "Bên cạnh các sơ đồ use case, sequence và class, project còn cần một lớp thiết kế dữ liệu để phục vụ lưu tiến trình chơi. Trong game có nhiều trạng thái thay đổi theo thời gian như vị trí người chơi, chỉ số sinh tồn, inventory, ngày hiện tại, nhiệm vụ, hội thoại đã kích hoạt và các cài đặt âm thanh/độ nhạy chuột. Nếu các trạng thái này chỉ tồn tại trong scene runtime, người chơi sẽ mất toàn bộ tiến trình khi thoát game hoặc chuyển scene.",
            "Normal",
        ),
        (
            "Thiết kế lưu trữ được tổ chức theo hướng gom dữ liệu vào các lớp có khả năng serialize. AllGameData đóng vai trò gói dữ liệu tổng, bên trong gồm dữ liệu người chơi, inventory, story, quest và dialogue. PlayerData lưu các thông tin liên quan trực tiếp đến nhân vật như vị trí, hướng nhìn, stamina và các chỉ số sinh tồn. Cách tổ chức này giúp SaveManager không phụ thuộc quá sâu vào từng component riêng lẻ mà có thể thu thập dữ liệu qua một điểm điều phối chung.",
            "Normal",
        ),
        ("Luồng lưu game được thiết kế theo các bước chính:", "Normal"),
        ("Người chơi chọn slot lưu hoặc hệ thống gọi yêu cầu lưu.", "List Number"),
        ("SaveManager đọc trạng thái hiện tại từ player, inventory, quest, dialogue và các service liên quan.", "List Number"),
        ("Dữ liệu được đóng gói vào AllGameData và ghi ra file tương ứng với slot.", "List Number"),
        ("UI slot cập nhật thông tin để người chơi biết slot nào đã có dữ liệu.", "List Number"),
        ("Luồng tải game được thiết kế theo chiều ngược lại:", "Normal"),
        ("Người chơi chọn slot có dữ liệu hợp lệ.", "List Number"),
        ("SaveManager đọc file lưu, kiểm tra dữ liệu và nạp scene cần thiết.", "List Number"),
        ("Sau khi scene sẵn sàng, hệ thống khôi phục vị trí người chơi, inventory, nhiệm vụ, hội thoại và các trạng thái liên quan.", "List Number"),
        ("Các UI như HUD nhiệm vụ, inventory và menu được đồng bộ lại để phản ánh dữ liệu vừa tải.", "List Number"),
        (
            "Trong phạm vi đồ án, thiết kế save/load ưu tiên chứng minh khả năng lưu tiến trình demo và khôi phục các dữ liệu quan trọng. Một số trạng thái scene phức tạp như toàn bộ object đã bị phá hủy, trạng thái animation chi tiết hoặc lịch sử tương tác nhỏ vẫn cần tiếp tục chuẩn hóa nếu phát triển thành sản phẩm hoàn chỉnh.",
            "Normal",
        ),
    ]:
        insert_before(marker, text, style)


def add_menu_settings_module(doc: Document):
    marker = doc.paragraphs[find_paragraph(doc, "4.2.2. Module điều khiển người chơi")]
    for text, style in [
        ("4.2.1.1. Bổ sung module save/load và cài đặt", "Heading 3"),
        (
            "Ngoài luồng vào game cơ bản, project còn có các script phục vụ lưu/tải game và cấu hình trải nghiệm người chơi. MainMenu xử lý các nút New Game, Load và Quit, đồng thời liên kết với âm thanh nền menu, chế độ toàn màn hình/cửa sổ, âm lượng và độ nhạy chuột. MenuManager điều khiển pause menu trong scene gameplay, khóa hoặc mở lại chuột và tạm dừng một số script điều khiển player khi menu đang hiển thị.",
            "Normal",
        ),
        (
            "SaveManager là thành phần trung tâm của nhóm save/load. Script này gom dữ liệu từ PlayerData, PlayerInventory, quest, story và dialogue vào AllGameData, sau đó ghi hoặc đọc dữ liệu theo từng slot. SaveSlot và LoadSlot chịu trách nhiệm ở tầng giao diện: SaveSlot hiển thị cảnh báo khi ghi đè dữ liệu cũ, còn LoadSlot chỉ cho phép tải khi slot tương ứng có file lưu hợp lệ.",
            "Normal",
        ),
        (
            "Các thiết lập menu được tách thêm qua MenuSettingsService và SubMenuSettingsController. Cách tách này giúp cùng một nhóm cài đặt có thể được dùng ở main menu hoặc sub menu mà không phải viết lại toàn bộ logic. Với đồ án, phần này cho thấy project không chỉ có gameplay trong scene chính mà còn có luồng trải nghiệm đầy đủ hơn: bắt đầu game, tạm dừng, chỉnh cài đặt, lưu và tải tiến trình.",
            "Normal",
        ),
    ]:
        insert_before(marker, text, style)


def add_evaluation(doc: Document):
    marker = doc.paragraphs[find_paragraph(doc, "4.4. Minh họa sản phẩm thực tế")]
    for text, style in [
        ("4.3.4. Đánh giá theo tiêu chí kỹ thuật", "Heading 3"),
        (
            "Sau khi kiểm thử các chức năng chính, project được đánh giá theo một số tiêu chí kỹ thuật thay vì chỉ dừng ở việc chức năng có chạy hay không. Các tiêu chí này giúp nhìn rõ mức độ hoàn thiện của sản phẩm demo và chỉ ra hướng cải tiến tiếp theo.",
            "Normal",
        ),
        ("Bảng 4- 3. Đánh giá sản phẩm theo tiêu chí kỹ thuật", "Caption"),
    ]:
        insert_before(marker, text, style)

    rows = [
        ["Tiêu chí", "Kết quả quan sát", "Nhận xét"],
        [
            "Tính hoàn chỉnh chức năng",
            "Các luồng chính như menu, di chuyển, tương tác, inventory, quest, dialogue, câu cá và nấu ăn đã hoạt động ở mức demo.",
            "Đáp ứng mục tiêu đồ án, nhưng cần thêm nội dung ending và nhiều tình huống phụ nếu phát triển thành game hoàn chỉnh.",
        ],
        [
            "Tính module hóa",
            "Mã nguồn được chia theo nhóm Player, Interact, Inventory, Quest, Dialogue, Menu, AudioManager và Rendering.",
            "Cách chia module giúp dễ trình bày, dễ sửa lỗi và thuận lợi khi mở rộng từng hệ thống riêng.",
        ],
        [
            "Tính mở rộng dữ liệu",
            "Item, quest và dialogue được tách thành các lớp dữ liệu/asset cấu hình ở nhiều vị trí.",
            "Giảm hard-code, nhưng vẫn cần chuẩn hóa thêm naming và reference giữa scene object với dữ liệu.",
        ],
        [
            "Trải nghiệm người chơi",
            "Game có HUD nhiệm vụ, prompt tương tác, inventory, dialogue UI, mini game và âm thanh phản hồi.",
            "Người chơi nhận được phản hồi khá rõ, nhưng UI runtime nên được chuẩn hóa thành prefab để ổn định hơn.",
        ],
        [
            "Khả năng bảo trì",
            "Các script chính đã tách vai trò tương đối rõ, tuy nhiên một số luồng còn phụ thuộc tên object hoặc PlayerPrefs.",
            "Cần bổ sung quy ước đặt tên, reference trực tiếp và test tự động để giảm lỗi khi mở rộng.",
        ],
    ]
    add_table_before(marker, rows)

    for text, style in [
        (
            "Từ bảng đánh giá trên có thể thấy project đã đạt mục tiêu xây dựng một bản game có thể chơi thử và có nhiều hệ thống liên kết. Điểm cần cải thiện lớn nhất không nằm ở một chức năng đơn lẻ mà ở việc chuẩn hóa dữ liệu, prefab UI, cơ chế save/load và quy trình kiểm thử để sản phẩm ổn định hơn khi quy mô nội dung tăng lên.",
            "Normal",
        ),
        ("4.3.5. Rủi ro kỹ thuật và biện pháp khắc phục", "Heading 3"),
        (
            "Trong quá trình phát triển game 3D, một số rủi ro kỹ thuật thường xuất hiện do nhiều hệ thống chạy đồng thời trong cùng scene. Project đã xử lý được các trường hợp cơ bản, nhưng vẫn cần ghi nhận các rủi ro để định hướng hoàn thiện.",
            "Normal",
        ),
        ("Rủi ro input bị xung đột khi mở dialogue, inventory, fishing hoặc cooking. Biện pháp khắc phục là tập trung hóa trạng thái khóa/mở input, chỉ cho phép một mode tương tác chính hoạt động tại một thời điểm.", "List Bullet"),
        ("Rủi ro dữ liệu quest và dialogue không đồng bộ khi chuyển ngày. Biện pháp khắc phục là quy định rõ service chịu trách nhiệm cập nhật ngày hiện tại, reset trạng thái và phát sự kiện cho HUD/UI.", "List Bullet"),
        ("Rủi ro object trong scene bị phụ thuộc tên. Biện pháp khắc phục là dùng reference serialize trong Inspector, tag/layer rõ ràng hoặc ScriptableObject chứa cấu hình liên kết.", "List Bullet"),
        ("Rủi ro UI runtime khó bảo trì khi số lượng màn hình tăng. Biện pháp khắc phục là chuyển các panel ổn định thành prefab, tách controller khỏi phần dựng giao diện và thống nhất style UI.", "List Bullet"),
        ("Rủi ro save/load thiếu trạng thái chi tiết. Biện pháp khắc phục là mở rộng AllGameData theo từng nhóm dữ liệu, thêm version cho save file và kiểm thử tải lại ở nhiều mốc ngày khác nhau.", "List Bullet"),
    ]:
        insert_before(marker, text, style)


def add_conclusion_detail(doc: Document):
    marker = doc.paragraphs[find_paragraph(doc, "TÀI LIỆU THAM KHẢO")]
    for text, style in [
        (
            "Nhìn chung, giá trị chính của đồ án là chứng minh khả năng kết hợp nhiều mảng kiến thức trong một sản phẩm cụ thể: lập trình hướng đối tượng với C#, mô hình component của Unity, xử lý input thời gian thực, tương tác vật lý trong không gian 3D, quản lý UI, dữ liệu cấu hình, âm thanh và tiến trình nhiệm vụ. Việc các module có thể phối hợp với nhau trong cùng một vòng chơi cho thấy project đã vượt qua mức mô phỏng đơn lẻ và tiến gần hơn tới một prototype game hoàn chỉnh.",
            "Normal",
        ),
        (
            "Nếu tiếp tục phát triển, ưu tiên đầu tiên nên là ổn định nền tảng kỹ thuật trước khi mở rộng nội dung. Cụ thể, cần hoàn thiện save/load, chuẩn hóa prefab UI, giảm phụ thuộc tên object, bổ sung kiểm thử theo kịch bản và hoàn thiện nội dung ngày 6. Sau đó có thể mở rộng hệ thống vật phẩm, công thức nấu ăn, sinh vật, sự kiện môi trường và cốt truyện để tăng chiều sâu trải nghiệm.",
            "Normal",
        ),
    ]:
        insert_before(marker, text, style)


def main() -> int:
    if len(sys.argv) != 3:
        print("Usage: docx_expand_report.py <input.docx> <output.docx>")
        return 2

    source = Path(sys.argv[1])
    target = Path(sys.argv[2])
    doc = Document(source)

    add_timeline(doc)
    add_save_design(doc)
    add_menu_settings_module(doc)
    add_evaluation(doc)
    add_conclusion_detail(doc)

    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(target)
    print(target)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
